"""Convierte un archivo Markdown a DOCX con formato de oposición.

Uso:
  python md_to_docx.py <ruta_md> [ruta_docx_salida]

Formato aplicado:
- Fuente: Times New Roman 12
- Interlineado: sencillo
- Márgenes: 2,5 cm
- Encabezados con jerarquía (H1, H2, H3)
- Conserva negritas, cursivas y listas
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_LINE_SPACING


def configurar_estilo_base(doc: Document) -> None:
    estilo = doc.styles["Normal"]
    estilo.font.name = "Times New Roman"
    estilo.font.size = Pt(12)
    parrafo_fmt = estilo.paragraph_format
    parrafo_fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    parrafo_fmt.space_before = Pt(0)
    parrafo_fmt.space_after = Pt(6)

    for nivel in (1, 2, 3, 4):
        try:
            h = doc.styles[f"Heading {nivel}"]
            h.font.name = "Times New Roman"
            h.font.color.rgb = RGBColor(0, 0, 0)
            h.font.bold = True
            if nivel == 1:
                h.font.size = Pt(16)
            elif nivel == 2:
                h.font.size = Pt(14)
            elif nivel == 3:
                h.font.size = Pt(12)
            else:
                h.font.size = Pt(12)
        except KeyError:
            pass

    for seccion in doc.sections:
        seccion.top_margin = Cm(2.5)
        seccion.bottom_margin = Cm(2.5)
        seccion.left_margin = Cm(2.5)
        seccion.right_margin = Cm(2.5)


PATRON_NEGRITA = re.compile(r"\*\*(.+?)\*\*")
PATRON_CURSIVA = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")
PATRON_CODIGO = re.compile(r"`([^`]+)`")
PATRON_ENLACE = re.compile(r"\[([^\]]+)\]\([^\)]+\)")


def anadir_texto_formateado(parrafo, texto: str) -> None:
    """Añade texto con negritas/cursivas/enlaces inline a un párrafo."""
    texto = PATRON_ENLACE.sub(r"\1", texto)
    texto = PATRON_CODIGO.sub(r"\1", texto)

    posicion = 0
    eventos = []
    for m in PATRON_NEGRITA.finditer(texto):
        eventos.append((m.start(), m.end(), "bold", m.group(1)))
    for m in PATRON_CURSIVA.finditer(texto):
        solapa = any(s <= m.start() < e for s, e, _, _ in eventos)
        if not solapa:
            eventos.append((m.start(), m.end(), "italic", m.group(1)))
    eventos.sort(key=lambda x: x[0])

    for inicio, fin, tipo, contenido in eventos:
        if inicio > posicion:
            parrafo.add_run(texto[posicion:inicio])
        run = parrafo.add_run(contenido)
        if tipo == "bold":
            run.bold = True
        elif tipo == "italic":
            run.italic = True
        posicion = fin
    if posicion < len(texto):
        parrafo.add_run(texto[posicion:])


def convertir(ruta_md: Path, ruta_docx: Path) -> None:
    contenido = ruta_md.read_text(encoding="utf-8")
    doc = Document()
    configurar_estilo_base(doc)

    lineas = contenido.splitlines()
    i = 0
    en_blockquote = False
    while i < len(lineas):
        linea = lineas[i].rstrip()

        if not linea:
            i += 1
            continue

        if linea.startswith("# "):
            doc.add_heading(linea[2:], level=1)
        elif linea.startswith("## "):
            doc.add_heading(linea[3:], level=2)
        elif linea.startswith("### "):
            doc.add_heading(linea[4:], level=3)
        elif linea.startswith("#### "):
            doc.add_heading(linea[5:], level=4)
        elif linea.startswith("---"):
            p = doc.add_paragraph()
            p.add_run("_" * 60)
        elif linea.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run("« ")
            run.italic = True
            anadir_texto_formateado(p, linea[2:])
        elif re.match(r"^\d+\.\s", linea):
            p = doc.add_paragraph(style="List Number")
            anadir_texto_formateado(p, re.sub(r"^\d+\.\s", "", linea))
        elif linea.startswith("- ") or linea.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            anadir_texto_formateado(p, linea[2:])
        elif linea.startswith("|"):
            filas = []
            while i < len(lineas) and lineas[i].startswith("|"):
                filas.append(lineas[i])
                i += 1
            if len(filas) >= 2:
                celdas_cabecera = [c.strip() for c in filas[0].split("|")[1:-1]]
                tabla = doc.add_table(rows=1, cols=len(celdas_cabecera))
                tabla.style = "Light Grid"
                for idx, c in enumerate(celdas_cabecera):
                    tabla.rows[0].cells[idx].text = c
                for fila_raw in filas[2:]:
                    valores = [c.strip() for c in fila_raw.split("|")[1:-1]]
                    fila = tabla.add_row()
                    for idx, v in enumerate(valores):
                        if idx < len(fila.cells):
                            fila.cells[idx].text = v
            continue
        else:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.5)
            anadir_texto_formateado(p, linea)

        i += 1

    doc.save(str(ruta_docx))


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python md_to_docx.py <ruta_md> [ruta_docx_salida]")
        sys.exit(1)

    ruta_md = Path(sys.argv[1]).resolve()
    if not ruta_md.exists():
        print(f"No existe: {ruta_md}")
        sys.exit(1)

    if len(sys.argv) >= 3:
        ruta_docx = Path(sys.argv[2]).resolve()
    else:
        ruta_docx = ruta_md.with_suffix(".docx")

    convertir(ruta_md, ruta_docx)
    print(f"Generado: {ruta_docx}")
